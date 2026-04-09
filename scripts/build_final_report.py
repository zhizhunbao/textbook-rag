"""Build Final Report following the teacher's exact template structure.

Framing: Ottawa Economic Development — GenAI Internal Research Tool

Template sections:
  Final Report (cover) -> Title Page -> Acknowledgements -> Executive Summary
  -> Table of Contents -> Introduction -> Data Set/Feature Engineering
  -> Architectural Design -> Program Design -> Recommendations -> Appendices
"""

from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUTPUT = Path(__file__).resolve().parent.parent / "Final_Report_Textbook_RAG_v2.docx"

# ── Design ──
ACCENT = RGBColor(47, 84, 150)
FONT = "Calibri"


def shading(cell, color_hex: str):
    tc_pr = cell._element.get_or_add_tcPr()
    elm = tc_pr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex})
    tc_pr.append(elm)


def styled_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        shading(c, "2F5496")
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = FONT
                r.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = FONT
            if ri % 2 == 1:
                shading(c, "D6E4F0")
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def p_normal(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = FONT
    r.bold = bold
    return p


def p_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = FONT
    return p


def bullets(doc, items):
    for item in items:
        p_bullet(doc, item)


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


# ============================================================
def build():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(2.54)
        sec.right_margin = Cm(2.54)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)

    # ════════════════════════════════════════════════════════════
    # Cover: "Final Report"
    # ════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    h = doc.add_heading("Final Report", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Date Issued table
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Date Issued:"
    t.rows[0].cells[1].text = "April 7, 2026"
    for c in t.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11)
                r.font.name = FONT

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # Title Page
    # ════════════════════════════════════════════════════════════
    heading(doc, "Title Page", 1)
    doc.add_paragraph()

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run("GenAI-Based Internal Research Tool\nOttawa EcDev RAG v2.0")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = ACCENT
    r.font.name = FONT

    doc.add_paragraph()

    for label, value in [
        ("Project Title:", "GenAI-Based Internal Research Tool — Ottawa EcDev RAG v2.0"),
        ("Client:", "City of Ottawa — Economic Development Department"),
        ("Student Names:", "Peng Wang (041107730), Hye Ran Yoo (041145212), Jiaxing Yi (041127158)"),
        ("Course:", "CST 8505 — AI Software Development Project"),
        ("Instructor:", "Hari"),
        ("Submission Date:", "April 7, 2026"),
    ]:
        p = doc.add_paragraph()
        rl = p.add_run(label + "  ")
        rl.bold = True
        rl.font.size = Pt(12)
        rl.font.name = FONT
        rv = p.add_run(value)
        rv.font.size = Pt(12)
        rv.font.name = FONT

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # Acknowledgements
    # ════════════════════════════════════════════════════════════
    heading(doc, "Acknowledgements", 1)
    bullets(doc, [
        "We would like to thank our instructor Hari for providing guidance and feedback throughout the project development lifecycle.",
        "We acknowledge the City of Ottawa Economic Development Department for providing the quarterly report data that served as the foundation for this project.",
        "We thank the open-source communities behind LlamaIndex, Payload CMS, MinerU, ChromaDB, and Ollama, whose tools and frameworks made this project possible.",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # Executive Summary
    # ════════════════════════════════════════════════════════════
    heading(doc, "Executive Summary", 1)
    bullets(doc, [
        "Problem: The City of Ottawa Economic Development (EcDev) Department produces quarterly reports covering employment, GDP, building permits, tourism, and business investment. Analysts must manually search 12 PDF reports (Q1 2022 – Q4 2024) to locate specific economic indicators and compare cross-quarter trends. Standalone LLMs hallucinate statistics without source verification, which is unacceptable for government decision-making.",
        "Method: We built a Retrieval-Augmented Generation (RAG) platform using a three-tier architecture — Next.js 15 + Payload CMS frontend, Python/FastAPI engine with LlamaIndex, and ChromaDB + PostgreSQL data layer.",
        "The engine uses hybrid retrieval (BM25 lexical + Vector semantic) fused with Reciprocal Rank Fusion (RRF, k=60), providing citation-aware answers with page-level source references so analysts can verify every data point.",
        "Additional capabilities include LLM-based question generation with auto-scoring (LLM-as-Judge), and a 5-dimensional evaluation framework (Faithfulness, Relevancy, Correctness, Context Relevancy, Answer Relevancy).",
        "Key Results: Hybrid retrieval outperforms single-strategy approaches; citation accuracy verified against real page references; question generation achieves avg. relevance ~4.1/5.0 and clarity ~4.3/5.0; ~70% of generated questions reach 'understanding' depth or above.",
        "Conclusion: A well-structured RAG pipeline with proper chunking, metadata injection, and hybrid retrieval can deliver reliable, source-grounded answers for government economic data analysis use cases.",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # Table of Contents
    # ════════════════════════════════════════════════════════════
    heading(doc, "Table of Contents", 1)
    bullets(doc, [
        "1. Introduction",
        "2. Data Set / Feature Engineering",
        "3. Architectural Design",
        "4. Program Design",
        "5. Recommendations",
        "6. Appendices",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # Introduction
    # ════════════════════════════════════════════════════════════
    heading(doc, "Introduction", 1)

    p_normal(doc, "Problem Statement", bold=True, size=13)
    bullets(doc, [
        "The City of Ottawa Economic Development Department publishes quarterly economic update reports covering key indicators such as employment, GDP growth, building permits, tourism, and business investment. From Q1 2022 through Q4 2024, 12 such reports have been produced.",
        "EcDev analysts need to find specific economic indicators across multiple quarters — for example, comparing year-over-year employment trends or identifying quarterly shifts in building permit values. Manual search through these PDF reports is time-consuming and error-prone.",
        "Existing keyword search tools fail for conceptual or cross-quarter queries where exact terms may not appear in the source text (e.g., \"How did Ottawa's labour market change from 2022 to 2024?\").",
        "Standalone LLMs can generate plausible but incorrect economic statistics (hallucination) without any source verification, making them unreliable for government policy and decision-making contexts.",
        "There is no automated way to generate review questions or audit queries from report content for internal assessment purposes.",
    ])

    p_normal(doc, "Project Goals", bold=True, size=13)
    bullets(doc, [
        "Build a hybrid retrieval system combining BM25 (lexical) and Vector (semantic) search with Reciprocal Rank Fusion (RRF) for accurate retrieval across Ottawa EcDev quarterly reports.",
        "Generate citation-aware answers with page-level source references that analysts can verify against original quarterly report content.",
        "Provide a content management system (Payload CMS) for managing the report library with automated PDF ingestion — upload a new quarterly report, and it is automatically parsed, chunked, and indexed.",
        "Support LLM-based question generation with auto-scoring for internal auditing and review purposes.",
        "Implement a 5-dimensional evaluation framework for measuring response quality (Faithfulness, Relevancy, Correctness, Context Relevancy, Answer Relevancy).",
        "Offer dual LLM support — local Ollama for development and Azure OpenAI (GPT-4o) for production — with zero code changes.",
    ])

    p_normal(doc, "Background and Context", bold=True, size=13)
    bullets(doc, [
        "Retrieval-Augmented Generation (RAG) is an AI architecture pattern that combines information retrieval with generative language models to reduce hallucination and enable source attribution.",
        "Our project builds on the LlamaIndex framework, which provides standardized abstractions for RAG components including retrievers, synthesizers, evaluators, and query engines.",
        "The system uses ChromaDB as the vector store for semantic search and MinerU for high-fidelity PDF parsing. The frontend is built with Next.js 15 and Payload CMS 3.x.",
        "The ultimate goal is migration to the Ottawa GenAI Research Assistant (RAG-Project) for production deployment within the City of Ottawa infrastructure.",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # Data Set / Feature Engineering
    # ════════════════════════════════════════════════════════════
    heading(doc, "Data Set/Feature Engineering", 1)

    p_normal(doc, "Dataset Description", bold=True, size=13)
    bullets(doc, [
        "The primary dataset consists of 12 quarterly economic update PDF reports published by the City of Ottawa Economic Development Department, spanning Q1 2022 through Q4 2024 (~7.4 MB total).",
    ])

    styled_table(doc,
        ["Quarter", "File Name", "Size"],
        [
            ["Q1 2022", "ed_update_q1_2022.pdf", "707 KB"],
            ["Q2 2022", "ed_update_q2_2022.pdf", "421 KB"],
            ["Q3 2022", "ed_update_q3_2022.pdf", "518 KB"],
            ["Q4 2022", "ed_update_q4_2022.pdf", "946 KB"],
            ["Q1 2023", "ed_update_q1_2023.pdf", "767 KB"],
            ["Q2 2023", "ed_update_q2_2023.pdf", "368 KB"],
            ["Q3 2023", "ed_update_q3_2023.pdf", "768 KB"],
            ["Q4 2023", "ed_update_q4_2023.pdf", "766 KB"],
            ["Q1 2024", "ed_update_q1_2024.pdf", "581 KB"],
            ["Q2 2024", "ed_update_q2_2024.pdf", "437 KB"],
            ["Q3 2024", "ed_update_q3_2024.pdf", "1,007 KB"],
            ["Q4 2024", "ed_update_q4_2024.pdf", "880 KB"],
        ],
        widths=[3, 6, 3],
    )
    bullets(doc, [
        "Report topics include: employment & labour market, GDP growth, building permits & construction, tourism & visitor spending, business investment & development, and quarterly/annual comparisons.",
    ])

    p_normal(doc, "Data Preprocessing", bold=True, size=13)
    bullets(doc, [
        "Stage 1 — PDF Parsing (MinerU): Each quarterly report PDF is converted into structured Markdown preserving headings, tables, figures, and page boundaries.",
        "Stage 2 — TOC Extraction: Regex-based extractor identifies table of contents and maps section boundaries to page ranges for section-level metadata.",
        "Stage 3 — Chunking: Parsed text is split into chunks using sentence-level splitting with a 512-token window and 64-token overlap.",
        "Stage 4 — Embedding: Chunks are embedded using HuggingFace all-MiniLM-L6-v2 model (384 dimensions) and stored in ChromaDB with full metadata.",
    ])

    p_normal(doc, "Feature Engineering", bold=True, size=13)
    bullets(doc, [
        "Each chunk is enriched with metadata fields that enable scoped retrieval:",
    ])
    styled_table(doc,
        ["Metadata Field", "Description", "Usage"],
        [
            ["book_id", "Unique identifier for the source report", "Report-scoped retrieval filtering"],
            ["book_title", "Report title (e.g., 'ed_update_q3_2023')", "Display in citations"],
            ["chapter", "Section name from structure mapping", "Section-level filtering"],
            ["page_idx", "Zero-indexed page number in original PDF", "Citation linking to exact page"],
            ["category", "Document category ('ecdev')", "Collection-level filtering"],
            ["heading_path", "Full heading hierarchy path", "Context enrichment"],
        ],
        widths=[4, 6, 6],
    )
    bullets(doc, [
        "These metadata fields enable MetadataFilters in ChromaDB, allowing the hybrid retriever to scope queries to specific reports, quarters, or page ranges, significantly improving retrieval precision for cross-quarter comparisons.",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # Architectural Design
    # ════════════════════════════════════════════════════════════
    heading(doc, "Architectural Design", 1)

    p_normal(doc, "End-to-End Solution Architecture", bold=True, size=13)
    bullets(doc, [
        "The system follows a three-tier modular architecture designed for separation of concerns and independent scalability.",
    ])

    p_normal(doc, "Tier 1 — Frontend Layer (Next.js 15 + Payload CMS 3.x):", bold=True)
    bullets(doc, [
        "Next.js 15 App Router serves the user-facing SPA with pages for Chat, Report Library, Question Generation, Evaluation, and Settings.",
        "Payload CMS 3.x provides an auto-generated Admin Panel with 14 collections: Books, Chapters, Chunks, IngestTasks, ChatSessions, ChatMessages, Queries, Questions, Evaluations, Llms, Prompts, PdfUploads, Media, Users.",
        "Auto-generated REST and GraphQL APIs for all collections eliminate manual API development.",
        "Role-based access control (admin / editor / reader) via JWT authentication.",
    ])

    p_normal(doc, "Tier 2 — Engine Layer (Python FastAPI):", bold=True)
    bullets(doc, [
        "FastAPI application with modular endpoints for query, ingestion, question generation, evaluation, and report management.",
        "Built on LlamaIndex core abstractions for standardized RAG components.",
    ])
    styled_table(doc,
        ["Module", "Purpose"],
        [
            ["readers/", "MinerUReader for PDF-to-Document conversion"],
            ["chunking/", "Section extraction and sentence-level text splitting"],
            ["ingestion/", "IngestionPipeline with transformations to ChromaDB"],
            ["retrievers/", "HybridRetriever: BM25 + Vector with RRF fusion"],
            ["response_synthesizers/", "CitationSynthesizer for source-attributed answers"],
            ["query_engine/", "RetrieverQueryEngine composing retriever + synthesizer"],
            ["question_gen/", "LLM-based question generation with auto-scoring"],
            ["evaluation/", "5-dimensional response evaluation"],
            ["llms/", "LLM resolver for Ollama / Azure OpenAI routing"],
        ],
        widths=[5.5, 10.5],
    )

    p_normal(doc, "Tier 3 — Data Layer:", bold=True)
    bullets(doc, [
        "ChromaDB: Persistent vector store with HNSW cosine similarity index and BM25 full-text index.",
        "PostgreSQL: Relational database managed by Payload CMS for all structured data.",
        "Ollama: Local LLM inference server (default: qwen3.5:4b).",
        "Azure OpenAI: Cloud LLM option (GPT-4o-mini) with automatic fallback to Ollama.",
    ])

    p_normal(doc, "Hybrid Retrieval Architecture", bold=True, size=13)
    bullets(doc, [
        "BM25 Retriever: Uses rank_bm25 over LlamaIndex docstore for keyword-based lexical matching. Effective for exact economic indicator queries (e.g., 'building permits Q3 2023').",
        "Vector Retriever: Uses ChromaDB HNSW index with cosine similarity for semantic matching. Effective for conceptual queries (e.g., 'What drove Ottawa employment growth?').",
        "QueryFusionRetriever (RRF): Combines both retrievers using Reciprocal Rank Fusion (k=60). Formula: RRF(d) = Σ 1/(k + rank_i(d)).",
        "MetadataFilters: Report-scoped filtering pushed down to ChromaDB native where clause to prevent cross-report contamination.",
        "Graceful degradation: Falls back to vector-only mode if ChromaDB collection is empty.",
    ])

    p_normal(doc, "LLM Integration", bold=True, size=13)
    bullets(doc, [
        "The LLM resolver implements priority-based routing: Azure OpenAI (Priority 1, if configured) → Ollama local (Priority 2, always available).",
        "Dual-mode design allows development with free local models and production with cloud models, zero code changes.",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # Program Design
    # ════════════════════════════════════════════════════════════
    heading(doc, "Program Design", 1)

    p_normal(doc, "Technology Stack", bold=True, size=13)
    styled_table(doc,
        ["Layer", "Technologies"],
        [
            ["Frontend", "Next.js 15 (App Router), TypeScript, Tailwind CSS, Payload CMS 3.x, React PDF"],
            ["Backend", "Python 3.12, FastAPI, LlamaIndex, MinerU, Loguru, uv (package manager)"],
            ["AI / ML", "Ollama (qwen3.5:4b), Azure OpenAI (GPT-4o-mini), HuggingFace all-MiniLM-L6-v2, ChromaDB"],
            ["Infrastructure", "PostgreSQL 15+, Docker Compose, JWT Auth, CORS"],
        ],
        widths=[3, 13],
    )

    p_normal(doc, "Code Structure", bold=True, size=13)
    styled_table(doc,
        ["Directory", "Purpose"],
        [
            ["engine_v2/api/", "FastAPI app, routes (query, ingest, books, questions, evaluation, llms)"],
            ["engine_v2/readers/", "MinerU PDF reader"],
            ["engine_v2/chunking/", "Section extraction + text splitting"],
            ["engine_v2/ingestion/", "IngestionPipeline to ChromaDB"],
            ["engine_v2/retrievers/", "HybridRetriever (BM25 + Vector → RRF)"],
            ["engine_v2/response_synthesizers/", "CitationSynthesizer"],
            ["engine_v2/question_gen/", "Question generation + scoring"],
            ["engine_v2/evaluation/", "5-dimensional evaluator"],
            ["engine_v2/llms/", "LLM resolver (Ollama / Azure)"],
            ["payload-v2/src/collections/", "14 Payload CMS collection definitions"],
            ["payload-v2/src/app/(frontend)/", "Next.js pages (chat, readers, engine, settings)"],
        ],
        widths=[6.5, 9.5],
    )

    p_normal(doc, "Key Algorithms and Custom Functions", bold=True, size=13)

    p_normal(doc, "1. Hybrid Retrieval with RRF Fusion:", bold=True)
    bullets(doc, [
        "HybridRetriever builds a QueryFusionRetriever combining BM25 + Vector retrievers.",
        "RRF formula: RRF(d) = Σ 1/(k + rank_i(d)), k=60. Produces a unified ranking benefiting from both lexical precision and semantic recall.",
    ])

    p_normal(doc, "2. LLM-as-Judge Scoring:", bold=True)
    bullets(doc, [
        "Two-phase approach: (1) generate questions from sampled chunks using structured JSON prompt, (2) score each question on relevance, clarity, and difficulty.",
        "Overall score = average(relevance, clarity).",
    ])

    p_normal(doc, "3. Question Depth Assessment:", bold=True)
    bullets(doc, [
        "QuestionDepthEvaluator extends LlamaIndex CorrectnessEvaluator with cognitive depth scoring (1-5 scale, Bloom's taxonomy).",
        "Labels: Surface (<2.5), Understanding (2.5-4.0), Synthesis (≥4.0).",
    ])

    p_normal(doc, "4. 5-Dimensional Response Evaluation:", bold=True)
    styled_table(doc,
        ["Evaluator", "Question Answered"],
        [
            ["FaithfulnessEvaluator", "Is the answer grounded in retrieved context?"],
            ["RelevancyEvaluator", "Is the retrieved context relevant to the query?"],
            ["CorrectnessEvaluator", "Does the answer match a reference answer?"],
            ["ContextRelevancyEvaluator", "Quality score of the retrieved contexts"],
            ["AnswerRelevancyEvaluator", "How well does the answer address the query?"],
        ],
        widths=[6, 10],
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # Recommendations
    # ════════════════════════════════════════════════════════════
    heading(doc, "Recommendations", 1)

    p_normal(doc, "Future Improvements", bold=True, size=13)

    p_normal(doc, "Short Term:", bold=True)
    bullets(doc, [
        "Docker Compose deployment for one-click full-stack startup.",
        "Streaming response generation for improved user experience during long answers.",
        "Batch ingestion via directory scanning to auto-detect and process new quarterly reports.",
        "Enhanced PDF viewer integration with bounding-box citation highlighting.",
    ])
    p_normal(doc, "Medium Term:", bold=True)
    bullets(doc, [
        "Azure AI Search as a 6th retrieval strategy for cloud-native semantic search.",
        "Migration to Ottawa GenAI Research Assistant (RAG-Project) for production deployment within City of Ottawa infrastructure.",
        "Entra ID authentication for enterprise deployment within City of Ottawa.",
        "Statistics Canada API integration for enriching responses with national economic context.",
        "Response caching for frequently asked queries to reduce LLM inference costs.",
    ])
    p_normal(doc, "Long Term:", bold=True)
    bullets(doc, [
        "Multi-modal retrieval supporting charts, infographics, and maps from quarterly reports.",
        "Conversational memory with context carryover across multiple turns.",
        "Fine-tuned embedding model trained on government economic data corpus.",
        "Cross-department expansion to Planning, Transit, and Tourism departments.",
    ])

    p_normal(doc, "Limitations and Ethical Considerations", bold=True, size=13)
    bullets(doc, [
        "PDF Parsing Quality: MinerU handles most formats well, but complex tables, charts, and economic infographics in EcDev reports may lose fidelity during parsing.",
        "LLM Evaluation Consistency: LLM-as-Judge scores can vary between runs; mitigated with structured prompts and batch averaging.",
        "Local LLM Performance: Smaller local models (qwen3.5:4b) produce lower quality compared to cloud models (GPT-4o), but offer privacy and cost benefits.",
        "BM25 Limitations: Requires at least one document in corpus; system falls back to vector-only mode on empty collections.",
        "Data Accuracy: Economic statistics in reports are published by the City of Ottawa. The RAG system retrieves and cites these numbers but does not independently verify them. Users should always cross-reference with official published data.",
        "Data Privacy: User queries are stored in PostgreSQL. Production deployment within City of Ottawa should implement data retention policies and comply with MFIPPA (Municipal Freedom of Information and Protection of Privacy Act).",
        "AI Transparency: Citation-aware answers promote transparency by allowing analysts to verify claims against original quarterly report pages.",
    ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # Appendices
    # ════════════════════════════════════════════════════════════
    heading(doc, "Appendices", 1)

    p_normal(doc, "Appendix A: API Endpoints", bold=True, size=13)
    styled_table(doc,
        ["Method", "Endpoint", "Purpose"],
        [
            ["POST", "/api/query", "Execute RAG query with optional report filters"],
            ["POST", "/api/ingest", "Trigger PDF ingestion pipeline"],
            ["GET", "/api/books", "List all reports with metadata"],
            ["POST", "/api/questions/generate", "Generate review questions"],
            ["POST", "/api/evaluation/evaluate", "Run 5-dimensional evaluation"],
            ["GET", "/api/llms", "List available LLM providers"],
            ["GET", "/api/health", "Health check"],
        ],
        widths=[2, 5.5, 8.5],
    )

    p_normal(doc, "Appendix B: Environment Configuration", bold=True, size=13)
    styled_table(doc,
        ["Variable", "Default", "Description"],
        [
            ["OLLAMA_BASE_URL", "http://127.0.0.1:11434", "Ollama server URL"],
            ["OLLAMA_MODEL", "qwen3.5:4b", "Local LLM model"],
            ["AZURE_OAI_ENDPOINT", "(optional)", "Azure OpenAI endpoint"],
            ["AZURE_OAI_KEY", "(optional)", "Azure OpenAI API key"],
            ["AZURE_OAI_DEPLOYMENT", "gpt-4o-mini", "Azure model deployment"],
            ["EMBEDDING_MODEL", "all-MiniLM-L6-v2", "HuggingFace embedding model"],
            ["CHROMA_PERSIST_DIR", "data/chroma_persist", "ChromaDB persistence dir"],
            ["PAYLOAD_URL", "http://localhost:3001", "Payload CMS URL"],
            ["TOP_K", "5", "Retrieval result count"],
        ],
        widths=[5, 4.5, 6.5],
    )

    p_normal(doc, "Appendix C: Project Repository", bold=True, size=13)
    bullets(doc, [
        "GitHub: github.com/Teegee0/textbook-rag",
        "Branch: main",
        "Python: 3.12 (managed with uv)",
        "Node.js: 20+ (for Payload CMS / Next.js)",
    ])

    # ── Save ──
    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
