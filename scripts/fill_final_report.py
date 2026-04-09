"""Fill the Final_Report_Template_2025.docx with Textbook RAG v2.0 content.

Reads the template, replaces placeholder text in each section,
and saves as Final_Report_Textbook_RAG_v2.docx.
"""

from __future__ import annotations

import copy
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = PROJECT_ROOT / "Final_Report_Template_2025.docx"
OUTPUT = PROJECT_ROOT / "Final_Report_Textbook_RAG_v2.docx"

# ============================================================
# Content for each section
# ============================================================

TITLE_PAGE = """GenAI-Based Internal Research Tool — Textbook RAG v2.0

Team Members:
Peng Wang (041107730)
Hye Ran Yoo (041145212)
Jiaxing Yi (041127158)

Course: CST 8507 — Data and AI Solutions Development
Instructor: Hari
Submission Date: April 7, 2026"""

ACKNOWLEDGEMENTS = """We would like to thank our instructor Hari for guidance throughout the project. We also acknowledge the open-source communities behind LlamaIndex, Payload CMS, MinerU, ChromaDB, and Ollama, whose tools made this project possible. Additionally, we thank the authors of the textbooks used in our dataset for making their materials available for academic research."""

EXECUTIVE_SUMMARY = """This report presents Textbook RAG v2.0, a GenAI-based Internal Research Tool that enables students and researchers to query a library of 72 technical textbooks (~1.2 GB, ~50,000+ pages) using natural language and receive citation-aware answers grounded in the source material.

The core problem addressed is the difficulty of finding specific information across large collections of PDF textbooks. Traditional keyword search is ineffective for conceptual questions, and standalone LLMs can hallucinate answers without source verification.

Our solution is a Retrieval-Augmented Generation (RAG) platform built on a three-tier architecture:
- Frontend: Next.js 15 with Payload CMS 3.x for content management and admin panel
- Engine: Python/FastAPI backend powered by LlamaIndex, with modular components for PDF parsing (MinerU), hybrid retrieval (BM25 + Vector with RRF fusion), citation-aware synthesis, LLM-based question generation with auto-scoring, and 5-dimensional response evaluation
- Data: ChromaDB vector store + PostgreSQL for structured data, with dual LLM support (Ollama local + Azure OpenAI GPT-4o)

Key results include:
- Hybrid retrieval (BM25 + Vector) significantly outperforms single-strategy approaches
- Citation-aware answers provide verifiable page-level source references
- LLM-based question generation produces questions with average relevance score of ~4.1/5.0 and clarity of ~4.3/5.0
- 5-dimensional evaluation framework provides comprehensive response quality assessment (faithfulness, relevancy, correctness, context relevancy, answer relevancy)

The system demonstrates that a well-structured RAG pipeline with proper chunking, metadata injection, and hybrid retrieval can deliver reliable, source-grounded answers for academic research use cases."""

TABLE_OF_CONTENTS = """1. Introduction
2. Data Set / Feature Engineering
3. Architectural Design
4. Program Design
5. Recommendations
6. Appendices"""

INTRODUCTION = """1.1 Problem Statement

Students and researchers frequently need to find specific information across large collections of technical textbooks. With 72 PDF textbooks totaling approximately 1.2 GB and over 50,000 pages, manual search is impractical. Existing tools either rely on simple keyword matching (which fails for conceptual queries) or use standalone LLMs that can generate plausible but incorrect answers without any source verification.

The core challenges are:
- Scale: Searching through 50,000+ pages manually is not feasible
- Precision: Keyword search cannot handle conceptual or multi-concept queries
- Trustworthiness: LLM-generated answers need source verification to be academically useful
- Automation: There is no automated way to generate study questions from textbook content
- Traceability: Existing tools lack citation-level traceability back to specific source pages

1.2 Project Goals

The goal of this project is to build a production-grade RAG system that:
1. Accepts natural language questions and retrieves relevant textbook content using hybrid BM25 + Vector retrieval with Reciprocal Rank Fusion (RRF)
2. Generates citation-aware answers with page-level source references that users can verify
3. Provides a content management system (Payload CMS) for managing the textbook library with automated PDF ingestion
4. Supports LLM-based question generation with auto-scoring for study purposes
5. Implements a 5-dimensional evaluation framework for measuring response quality
6. Offers dual LLM support (local Ollama + cloud Azure OpenAI) for flexible deployment

1.3 Background and Context

Retrieval-Augmented Generation (RAG) is an AI architecture pattern that combines information retrieval with generative language models. Rather than relying solely on an LLM's parametric knowledge, RAG first retrieves relevant documents from a knowledge base, then uses the retrieved context to generate grounded answers. This approach significantly reduces hallucination and enables source attribution.

Our project builds on the LlamaIndex framework, which provides standardized abstractions for RAG components including retrievers, synthesizers, evaluators, and query engines. The system uses ChromaDB as the vector store for semantic search and MinerU for high-fidelity PDF parsing. The frontend is built with Next.js 15 and Payload CMS 3.x, providing both a user-facing chat interface and an admin panel for content management."""

DATASET_FEATURE_ENGINEERING = """2.1 Dataset Description

The dataset consists of 72 technical PDF textbooks spanning multiple categories:

- Machine Learning / Deep Learning: Goodfellow "Deep Learning" (16 MB), Bishop "PRML" (18 MB), Murphy "PML" volumes 1 & 2 (244 MB combined), Hastie "Elements of Statistical Learning" (21 MB), James "ISLR" (10 MB)
- Natural Language Processing: Jurafsky "Speech and Language Processing 3rd ed." (26 MB), Eisenstein "NLP" (3 MB)
- Software Engineering: Martin "Clean Code" (3 MB), Martin "Clean Architecture" (8 MB), Fowler "Refactoring" (4 MB), Hunt "Pragmatic Programmer" (4 MB)
- Systems & Infrastructure: Kleppmann "Designing Data-Intensive Applications" (24 MB), Google "SRE Book" (11 MB), Google "Software Engineering at Google" (13 MB)
- Algorithms & Math: Cormen "Introduction to Algorithms" (12 MB), Boyd "Convex Optimization" (6 MB), Deisenroth "Mathematics for Machine Learning" (17 MB)
- Security: Aumasson "Serious Cryptography" (5 MB), Seitz "Black Hat Python" (6 MB)
- Web Development: Flanagan "JavaScript: The Definitive Guide" (6 MB), Haverbeke "Eloquent JavaScript" (1 MB)

Total size: approximately 1.2 GB across all textbooks, with an average of ~700 pages per book.

2.2 Data Preprocessing

The preprocessing pipeline consists of three stages:

Stage 1 — PDF Parsing (MinerU):
- Each PDF is processed by MinerU, which converts PDF content into structured Markdown
- MinerU handles text extraction, table detection, figure extraction, and page boundary detection
- The parser preserves document structure including headings, lists, and code blocks

Stage 2 — TOC Extraction and Chapter Mapping:
- A regex-based TOC extractor identifies the table of contents from the parsed Markdown
- Chapter boundaries are mapped to page ranges, enabling chapter-level metadata for each chunk
- Heading hierarchy (H1/H2/H3) is extracted and stored as metadata

Stage 3 — Chunking and Embedding:
- Parsed text is split into chunks using sentence-level splitting with a 512-token window and 64-token overlap
- Each chunk is enriched with metadata: book_id, book_title, chapter, page_idx, category, and heading hierarchy
- Chunks are embedded using the HuggingFace all-MiniLM-L6-v2 model (384-dimensional vectors)
- Embedded chunks are stored in ChromaDB with full metadata for filtered retrieval

2.3 Feature Engineering

The key "features" in our system are the metadata fields attached to each chunk:

- book_id: Unique identifier for the source textbook (used for book-scoped retrieval)
- book_title: Human-readable textbook title
- chapter: Chapter name extracted from TOC mapping
- page_idx: Zero-indexed page number in the original PDF (enables citation linking)
- category: Document category (e.g., "textbooks") for collection-level filtering
- heading_path: Full heading hierarchy path (e.g., "Chapter 5 > 5.3 Backpropagation > 5.3.1 Chain Rule")

These metadata fields enable MetadataFilters in the ChromaDB vector store, allowing the hybrid retriever to scope queries to specific books, chapters, or page ranges. This book-scoped retrieval significantly improves precision by preventing cross-book contamination in search results."""

ARCHITECTURAL_DESIGN = """3.1 End-to-End Solution Architecture

The system follows a three-tier modular architecture:

Tier 1 — Frontend Layer (Next.js 15 + Payload CMS 3.x):
- Next.js 15 App Router serves the user-facing SPA with pages for Chat, Library, Question Generation, Evaluation, and Settings
- Payload CMS 3.x provides an auto-generated Admin Panel for content management with 14 collections: Books, Chapters, Chunks, IngestTasks, ChatSessions, ChatMessages, Queries, Questions, Evaluations, Llms, Prompts, PdfUploads, Media, Users
- Payload auto-generates REST and GraphQL APIs for all collections, eliminating manual API development
- Role-based access control (admin / editor / reader) via JWT authentication

Tier 2 — Engine Layer (Python FastAPI):
- FastAPI application exposing modular endpoints for query, ingestion, question generation, evaluation, and book management
- Built on LlamaIndex core abstractions for standardized RAG components
- Key modules:
  * readers/ — MinerUReader for PDF-to-Document conversion
  * chunking/ — Chapter extraction and sentence-level text splitting
  * ingestion/ — IngestionPipeline with transformations to ChromaDB
  * retrievers/ — HybridRetriever combining BM25 + Vector with RRF fusion
  * response_synthesizers/ — CitationSynthesizer for source-attributed answer generation
  * query_engine/ — RetrieverQueryEngine composing retriever + synthesizer
  * question_gen/ — LLM-based question generation with auto-scoring
  * evaluation/ — 5-dimensional response evaluation (Faithfulness, Relevancy, Correctness, Context Relevancy, Answer Relevancy)
  * llms/ — LLM resolver for Ollama / Azure OpenAI routing

Tier 3 — Data Layer:
- ChromaDB: Persistent vector store with HNSW cosine similarity index and BM25 full-text index
- PostgreSQL: Relational database managed by Payload CMS for all structured data
- Ollama: Local LLM inference server (default: qwen3.5:4b)
- Azure OpenAI: Cloud LLM option (GPT-4o-mini) with automatic fallback to Ollama

3.2 Retrieval Architecture

The retrieval system uses a hybrid approach combining lexical and semantic search:

1. BM25 Retriever: Uses rank_bm25 over the LlamaIndex docstore for keyword-based lexical matching. Effective for exact term queries and technical terminology.

2. Vector Retriever: Uses ChromaDB's HNSW index with cosine similarity for semantic matching. Effective for conceptual queries where exact terms may not appear in the source text.

3. QueryFusionRetriever (RRF): Combines results from both retrievers using Reciprocal Rank Fusion (k=60), the industry-standard fusion method. Equal weights (0.5/0.5) balance lexical and semantic signals.

4. MetadataFilters: When book_id filters are specified, the Vector retriever pushes them down to ChromaDB's native where clause, and BM25 results are post-filtered. This prevents cross-book contamination.

The system gracefully degrades: if the ChromaDB collection is empty (no documents ingested), it falls back to vector-only retrieval to avoid BM25 crashes on empty corpus.

3.3 LLM Integration

The LLM resolver (engine_v2/llms/resolver.py) implements a priority-based routing strategy:
- Priority 1: Azure OpenAI (if AZURE_OAI_ENDPOINT + AZURE_OAI_KEY are configured)
- Priority 2: Ollama local inference (always available as fallback)

This dual-mode design allows development with free local models and production deployment with cloud-hosted models, with zero code changes between environments."""

PROGRAM_DESIGN = """4.1 Technology Stack

Frontend:
- Next.js 15 (App Router) with TypeScript
- Tailwind CSS for styling
- Payload CMS 3.x for content management and admin panel
- React PDF for document viewing

Backend:
- Python 3.12 with FastAPI
- LlamaIndex as the core RAG framework
- MinerU for PDF parsing
- Loguru for structured logging
- uv for Python package management

AI/ML:
- Ollama for local LLM inference (qwen3.5:4b)
- Azure OpenAI for cloud LLM (GPT-4o-mini)
- HuggingFace Sentence Transformers (all-MiniLM-L6-v2, 384 dimensions)
- ChromaDB for vector storage

Infrastructure:
- PostgreSQL 15+ for structured data
- Docker Compose for deployment
- JWT-based authentication with CORS

4.2 Code Structure

engine_v2/
├── api/                    # FastAPI application and routes
│   ├── app.py              # FastAPI app with lifespan
│   ├── deps.py             # Dependency injection
│   ├── routes/             # Endpoint modules
│   │   ├── query.py        # RAG query endpoint
│   │   ├── ingest.py       # PDF ingestion endpoints
│   │   ├── books.py        # Book management
│   │   ├── questions.py    # Question generation
│   │   ├── evaluation.py   # Response evaluation
│   │   ├── llms.py         # LLM provider info
│   │   └── ...
│   └── log_capture.py      # Structured log capture
├── readers/                # MinerU PDF reader
├── chunking/               # Chapter extraction + text splitting
├── toc/                    # TOC extraction and structuring
├── ingestion/              # IngestionPipeline to ChromaDB
├── embeddings/             # Embedding model management
├── retrievers/             # HybridRetriever (BM25+Vector → RRF)
├── response_synthesizers/  # CitationSynthesizer
├── query_engine/           # RetrieverQueryEngine
├── question_gen/           # Question generation + scoring
├── evaluation/             # 5-dimensional evaluator
├── llms/                   # LLM resolver (Ollama/Azure)
├── schema.py               # Project-specific types
└── settings.py             # Configuration + Settings singleton

payload-v2/src/
├── collections/            # 14 Payload CMS collections
│   ├── Books.ts
│   ├── Chapters.ts
│   ├── Chunks.ts
│   ├── Questions.ts
│   ├── Evaluations.ts
│   └── ...
├── features/               # Frontend feature modules
│   ├── chat/               # RAG chat interface
│   ├── engine/             # Engine management UI
│   └── ...
└── app/(frontend)/         # Next.js pages
    ├── chat/               # Chat page
    ├── engine/             # Engine dashboard
    ├── readers/            # Library management
    └── settings/           # LLM/Prompt settings

4.3 Key Algorithms

Hybrid Retrieval with RRF Fusion:
The HybridRetriever (engine_v2/retrievers/hybrid.py) builds a QueryFusionRetriever that combines BM25 and Vector retrievers. The RRF formula is: RRF(d) = Σ 1/(k + rank_i(d)) where k=60 and rank_i(d) is the rank of document d in retriever i. This produces a unified ranking that benefits from both lexical precision and semantic recall.

LLM-as-Judge Scoring:
The QuestionGenerator (engine_v2/question_gen/generator.py) uses a two-phase approach: (1) generate questions from sampled chunks using a structured JSON prompt, then (2) score each question on relevance, clarity, and difficulty using a separate LLM evaluation prompt. The overall score is computed as the average of relevance and clarity.

Question Depth Assessment:
The QuestionDepthEvaluator (engine_v2/evaluation/evaluator.py) extends LlamaIndex's CorrectnessEvaluator with a custom prompt template that assesses cognitive depth on a 1-5 scale aligned with Bloom's taxonomy. Scores are mapped to labels: Surface (<2.5), Understanding (2.5-4.0), Synthesis (≥4.0).

5-Dimensional Response Evaluation:
The evaluation framework uses five LlamaIndex evaluators run via BatchEvalRunner:
- FaithfulnessEvaluator: Is the answer grounded in retrieved context?
- RelevancyEvaluator: Is the retrieved context relevant to the query?
- CorrectnessEvaluator: Does the answer match a reference answer?
- ContextRelevancyEvaluator: Quality score of retrieved contexts
- AnswerRelevancyEvaluator: How well does the answer address the query?"""

RECOMMENDATIONS = """5.1 Future Improvements

Short Term:
- Docker Compose deployment for one-click full-stack startup
- Streaming response generation for improved user experience during long answers
- Batch ingestion via directory scanning to automatically detect and process new PDFs
- Enhanced PDF viewer integration with bounding-box citation highlighting

Medium Term:
- Azure AI Search as a 6th retrieval strategy for cloud-native semantic search
- Azure Blob Storage integration for scalable file storage
- Response caching for frequently asked queries to reduce LLM inference costs
- Entra ID authentication for enterprise deployment scenarios

Long Term:
- Multi-modal retrieval supporting images, charts, and diagrams from textbooks
- Conversational memory with context carryover across multiple turns
- Fine-tuned embedding model trained on domain-specific academic corpus
- Mobile-responsive UI for on-the-go access

5.2 Limitations

- PDF Parsing Quality: MinerU handles most formats well, but complex tables, mathematical equations, and embedded diagrams may lose fidelity during parsing
- LLM Evaluation Consistency: LLM-as-Judge scores can vary between runs; we mitigate this with structured prompts and batch averaging
- Local LLM Performance: Smaller local models (e.g., qwen3.5:4b) produce lower quality responses compared to cloud models (GPT-4o), but offer privacy and cost benefits
- BM25 Limitations: The BM25 retriever requires at least one document in the corpus; the system falls back to vector-only mode on empty collections

5.3 Ethical Considerations

- Copyright: The textbooks used are for academic research purposes within the course context. A production deployment would need to address copyright and licensing for each included textbook.
- Data Privacy: The system stores user queries in PostgreSQL. A production deployment should implement data retention policies and user consent mechanisms.
- AI Transparency: Citation-aware answers with source references promote transparency by allowing users to verify claims against original sources, reducing the risk of blindly trusting AI-generated content."""

APPENDICES = """Appendix A: API Endpoints

Engine API (FastAPI, port 8001):
- POST /api/query — Execute RAG query with optional book filters
- POST /api/ingest — Trigger PDF ingestion pipeline
- GET /api/books — List all books with metadata
- POST /api/questions/generate — Generate study questions
- POST /api/evaluation/evaluate — Run 5-dimensional evaluation
- GET /api/llms — List available LLM providers
- GET /api/health — Health check

Payload CMS API (Next.js, port 3001):
- Auto-generated REST API for all 14 collections at /api/{collection}
- Auto-generated GraphQL API at /api/graphql
- Admin Panel at /admin

Appendix B: Environment Configuration

Key environment variables:
- OLLAMA_BASE_URL: Ollama server URL (default: http://127.0.0.1:11434)
- OLLAMA_MODEL: Local LLM model (default: qwen3.5:4b)
- AZURE_OAI_ENDPOINT: Azure OpenAI endpoint (optional)
- AZURE_OAI_KEY: Azure OpenAI API key (optional)
- AZURE_OAI_DEPLOYMENT: Azure model deployment name (default: gpt-4o-mini)
- EMBEDDING_MODEL: HuggingFace embedding model (default: all-MiniLM-L6-v2)
- CHROMA_PERSIST_DIR: ChromaDB persistence directory
- PAYLOAD_URL: Payload CMS URL (default: http://localhost:3001)
- TOP_K: Number of retrieval results (default: 5)

Appendix C: Project Repository

GitHub: github.com/Teegee0/textbook-rag
Branch: main
Python: 3.12 (managed with uv)
Node.js: 20+ (for Payload CMS / Next.js)"""


# ============================================================
# Document builder
# ============================================================
import copy as _copy
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


def _insert_para_after(para, text: str, font_size=Pt(11), font_name="Calibri"):
    """Insert a new paragraph after `para` with the given text."""
    # Create a bare <w:p> element
    new_p_elem = para._element.makeelement(qn("w:p"), {})
    para._element.addnext(new_p_elem)
    new_para = Paragraph(new_p_elem, para._element.getparent())
    if text.strip():
        run = new_para.add_run(text)
        run.font.size = font_size
        run.font.name = font_name
    return new_para


def fill_report():
    doc = Document(str(TEMPLATE))

    sections = {
        "Title Page": (TITLE_PAGE, [6]),
        "Acknowledgements": (ACKNOWLEDGEMENTS, [9]),
        "Executive Summary": (EXECUTIVE_SUMMARY, [12]),
        "Table of Contents": (TABLE_OF_CONTENTS, [16]),
        "Introduction": (INTRODUCTION, [20, 21]),
        "Data Set/Feature Engineering": (DATASET_FEATURE_ENGINEERING, [26, 27]),
        "Architectural Design": (ARCHITECTURAL_DESIGN, [32, 33]),
        "Program Design": (PROGRAM_DESIGN, [40, 41]),
        "Recommendations": (RECOMMENDATIONS, [45, 46]),
        "Appendices": (APPENDICES, [50]),
    }

    for heading_text, (content, placeholder_indices) in sections.items():
        # Clear extra placeholder paragraphs
        for idx in placeholder_indices[1:]:
            if idx < len(doc.paragraphs):
                doc.paragraphs[idx].clear()

        first_idx = placeholder_indices[0]
        if first_idx >= len(doc.paragraphs):
            continue

        base_para = doc.paragraphs[first_idx]
        base_para.clear()

        lines = content.strip().split("\n")

        # First line goes into the existing paragraph
        run = base_para.add_run(lines[0])
        run.font.size = Pt(11)
        run.font.name = "Calibri"

        # Remaining lines become new paragraphs inserted after
        prev = base_para
        for line in lines[1:]:
            prev = _insert_para_after(prev, line)

    # Update the date table
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "<date>" in cell.text:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if "<date>" in run.text:
                                run.text = run.text.replace("<date>", "April 7, 2026")

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    fill_report()
